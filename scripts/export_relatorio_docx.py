# -*- coding: utf-8 -*-
"""
Exporta o relatório em Markdown deste projeto para .docx.

Entrada por defeito: `relatorio_investimentos.md` (raiz do repositório).
Saída por defeito: `relatorio_investimentos.docx` (mesma raiz).

Imagens: use no .md `![legenda](figures/nome.png)` com caminhos relativos
à pasta do ficheiro .md (ex.: `figures/fig01_evolucao.png`). Linhas
`*[Inserir figura: ...]*` são ignoradas no DOCX (marcadores de trabalho).
Se `figures/` não existir, corra `python scripts/gerar_figuras_relatorio.py`
antes do export (ou use `--gerar-figuras` neste script).

Inclui:
  - Margens (por defeito: esquerda 3 cm; demais 2 cm)
  - Recuo de primeira linha 1,0 cm no texto corrido
  - Títulos (níveis 1–3) alinhados à esquerda
  - Fonte Times New Roman ou Arial 12 pt (--font)
  - Citações longas (bloco Markdown > ) com recuo ~4 cm
  - Secção \"## Referências\" (ou \"## 5. Referências\", etc.): sangria pendente nas entradas bibliográficas
  - Espaçamento entre linhas ~1,15 no corpo; figuras ~10 cm de largura

Não gera: sumário automático.

Rodapé: numeração centrada (corpo reinicia em 1 após capa IDP, se existir).

**Capa IDP (opcional):** se existir `relatorio_cabecalho.json` na raiz do projeto
(ou caminho passado com `--cabecalho`), o DOCX abre com cabeçalho institucional
(disciplina, curso, docente, entrega, etc.). O campo `titulo_capa` (ex.: PROJETO FINAL)
é opcional; se estiver vazio ou ausente, essa linha não aparece. O corpo do relatório segue na **mesma secção seguinte**, **sem salto de página obrigatório** após o nome do aluno (o texto começa a seguir na página da capa se couber). Use `--no-cabecalho` para desativar.

Uso (na raiz do projeto):
    pip install python-docx
    python scripts/gerar_figuras_relatorio.py
    python scripts/export_relatorio_docx.py
    python scripts/export_relatorio_docx.py --gerar-figuras
    python scripts/export_relatorio_docx.py -i relatorio_investimentos.md -o output/relatorio.docx

Requer: python-docx
"""
from __future__ import annotations

import argparse
import json
import re
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

ROOT = Path(__file__).resolve().parent.parent
MD_PATH = ROOT / "relatorio_investimentos.md"
OUT_PATH = ROOT / "relatorio_investimentos.docx"

FONT_TIMES = "Times New Roman"
FONT_ARIAL = "Arial"

MARGIN_LEFT_CM = 3.0
MARGIN_OTHER_CM = 2.0
FIRST_LINE_CM = 1.0
QUOTE_INDENT_CM = 4.0
REF_HANGING_CM = 1.25


def set_run_font(run, size_pt: int = 12, font_name: str = FONT_TIMES) -> None:
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    try:
        r = run._element
        rPr = r.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn("w:ascii"), font_name)
        rFonts.set(qn("w:hAnsi"), font_name)
        rFonts.set(qn("w:eastAsia"), font_name)
        rFonts.set(qn("w:cs"), font_name)
    except Exception:
        pass


def _add_page_number_field_to_paragraph(p, font_name: str, size_pt: int = 11) -> None:
    """
    Numeração no rodapé: ``w:fldSimple`` com instrução PAGE (evita ``<w:r/>`` vazio
    que o ``paragraph.text = ''`` deixava e que impedia o Word de mostrar o campo).
    """
    p_el = p._element
    for child in list(p_el):
        if child.tag != qn("w:pPr"):
            p_el.remove(child)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_pr = p_el.find(qn("w:pPr"))
    if p_pr is None:
        p_pr = OxmlElement("w:pPr")
        p_el.insert(0, p_pr)
    if p_pr.find(qn("w:jc")) is None:
        jc = OxmlElement("w:jc")
        jc.set(qn("w:val"), "center")
        p_pr.append(jc)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), " PAGE \\* MERGEFORMAT ")
    r = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), font_name)
    r_fonts.set(qn("w:hAnsi"), font_name)
    r_fonts.set(qn("w:eastAsia"), font_name)
    r_fonts.set(qn("w:cs"), font_name)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size_pt * 2)))
    sz_cs = OxmlElement("w:szCs")
    sz_cs.set(qn("w:val"), str(int(size_pt * 2)))
    r_pr.extend([r_fonts, sz, sz_cs])
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = "\u00a0"
    r.append(r_pr)
    r.append(t)
    fld.append(r)
    p_el.append(fld)


def _enable_update_fields_on_open(doc: Document) -> None:
    """Pedido ao Word para atualizar campos (incl. PAGE) ao abrir o documento."""
    el = doc.settings.element
    for child in list(el):
        if child.tag == qn("w:updateFields"):
            el.remove(child)
    uf = OxmlElement("w:updateFields")
    uf.set(qn("w:val"), "true")
    el.append(uf)


def _section_set_page_number_start(section, start: int) -> None:
    """w:pgNumType na secção (reinício da numeração visível)."""
    sect_pr = section._sectPr
    for child in list(sect_pr):
        if child.tag == qn("w:pgNumType"):
            sect_pr.remove(child)
    pg = OxmlElement("w:pgNumType")
    pg.set(qn("w:start"), str(int(start)))
    sect_pr.append(pg)


def _setup_section_footer_page_number(section, font_name: str, *, restart_from: int | None) -> None:
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    _add_page_number_field_to_paragraph(fp, font_name)
    if restart_from is not None:
        _section_set_page_number_start(section, restart_from)


def _merge_empty_section_break_paragraph_into_previous(doc: Document) -> None:
    """
    O ``add_section`` do python-docx acrescenta um ``<w:p>`` vazio que só traz
    ``w:sectPr`` (quebra de secção). Em muitas versões do Word isso rende uma
    página quase em branco antes do corpo. Movemos o ``w:sectPr`` para o último
    parágrafo real (fim da capa) e removemos o parágrafo vazio.
    """
    if len(doc.paragraphs) < 2:
        return
    p_last = doc.paragraphs[-1]
    p_prev = doc.paragraphs[-2]
    el_last = p_last._element
    ppr_last = el_last.find(qn("w:pPr"))
    if ppr_last is None:
        return
    sect = ppr_last.find(qn("w:sectPr"))
    if sect is None:
        return
    ppr_last.remove(sect)
    if len(ppr_last) == 0:
        el_last.remove(ppr_last)
    el_prev = p_prev._element
    ppr_prev = el_prev.find(qn("w:pPr"))
    if ppr_prev is None:
        ppr_prev = OxmlElement("w:pPr")
        el_prev.insert(0, ppr_prev)
    ppr_prev.append(sect)
    parent = el_last.getparent()
    if parent is not None:
        parent.remove(el_last)


def apply_section_margins(section, left_cm: float, other_cm: float) -> None:
    section.left_margin = Cm(left_cm)
    section.right_margin = Cm(other_cm)
    section.top_margin = Cm(other_cm)
    section.bottom_margin = Cm(other_cm)


def format_paragraph_body(
    p,
    font_name: str,
    first_line: bool = True,
) -> None:
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.first_line_indent = Cm(FIRST_LINE_CM) if first_line else Cm(0)
    for run in p.runs:
        if run.font.name and "Courier" in (run.font.name or ""):
            continue
        set_run_font(run, 12, font_name)


def format_paragraph_blockquote(p, font_name: str) -> None:
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(QUOTE_INDENT_CM)
    p.paragraph_format.right_indent = Cm(QUOTE_INDENT_CM)
    p.paragraph_format.first_line_indent = Cm(0)
    for run in p.runs:
        set_run_font(run, 11, font_name)


def format_paragraph_reference_entry(p, font_name: str) -> None:
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(REF_HANGING_CM)
    p.paragraph_format.first_line_indent = Cm(-REF_HANGING_CM)
    for run in p.runs:
        set_run_font(run, 12, font_name)


def format_paragraph_caption(p, font_name: str) -> None:
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.first_line_indent = Cm(0)
    for run in p.runs:
        set_run_font(run, 12, font_name)


def format_list_item(p, font_name: str) -> None:
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.first_line_indent = Cm(0)
    for run in p.runs:
        set_run_font(run, 12, font_name)


def format_heading_left(p, font_name: str, size_pt: int) -> None:
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(3)
    # O estilo «Heading 1» do modelo base do Word costuma ter «quebra antes»,
    # o que deixa a 1.ª página do corpo quase em branco após a capa.
    p.paragraph_format.page_break_before = False
    for run in p.runs:
        set_run_font(run, size_pt, font_name)


def preprocess_links(text: str) -> str:
    return re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", text)


def strip_table_cell_markdown(text: str) -> str:
    """
    Células de tabela Markdown são copiadas com ``cell.text`` sem passar pelo
    interpretador de negrito/itálico/código — ``**termo**`` ou `` `ficheiro.csv` ``
    apareciam literais no Word.
    Remove marcação comum; preserva o texto visível.
    """
    t = (text or "").strip()
    for _ in range(16):
        n = re.sub(r"`([^`]*)`", r"\1", t)
        if n == t:
            break
        t = n
    for _ in range(4):
        n = re.sub(r"\*\*\*([^*]+)\*\*\*", r"\1", t)
        if n == t:
            break
        t = n
    for _ in range(8):
        n = re.sub(r"\*\*([^*]+)\*\*", r"\1", t)
        if n == t:
            break
        t = n
    for _ in range(8):
        n = re.sub(r"(?<!\*)\*([^*]+)\*(?!\*)", r"\1", t)
        if n == t:
            break
        t = n
    return t


def _emit_run(
    paragraph,
    chunk: str,
    size_pt: int,
    font_name: str,
    *,
    bold: bool = False,
    italic: bool = False,
    code: bool = False,
) -> None:
    if not chunk:
        return
    r = paragraph.add_run(chunk)
    if code:
        r.font.name = "Courier New"
        r.font.size = Pt(size_pt - 1)
    r.bold = bold
    r.italic = italic
    if not code:
        set_run_font(r, size_pt, font_name)
    elif bold or italic:
        # código com contexto raro: aplica só peso se necessário
        r.bold = bold
        r.italic = italic


def add_inline_runs(
    paragraph,
    text: str,
    size_pt: int,
    font_name: str,
    *,
    bold: bool = False,
    italic: bool = False,
) -> None:
    """
    Markdown inline: `` `código` ``, **negrito**, *itálico* (inclui *aninhado* dentro de **negrito**).
    Ordem: links, depois `, depois **, depois *.
    """
    text = preprocess_links(text)
    i = 0
    n = len(text)
    while i < n:
        if text[i] == "`":
            j = text.find("`", i + 1)
            if j == -1:
                _emit_run(paragraph, text[i : i + 1], size_pt, font_name, bold=bold, italic=italic)
                i += 1
                continue
            chunk = text[i + 1 : j]
            _emit_run(
                paragraph,
                chunk,
                size_pt,
                font_name,
                bold=bold,
                italic=italic,
                code=True,
            )
            i = j + 1
            continue

        if text.startswith("**", i):
            j = text.find("**", i + 2)
            if j == -1:
                _emit_run(paragraph, text[i : i + 2], size_pt, font_name, bold=bold, italic=italic)
                i += 2
                continue
            inner = text[i + 2 : j]
            add_inline_runs(
                paragraph,
                inner,
                size_pt,
                font_name,
                bold=True,
                italic=italic,
            )
            i = j + 2
            continue

        if text[i] == "*":
            j = text.find("*", i + 1)
            if j == -1:
                _emit_run(paragraph, text[i : i + 1], size_pt, font_name, bold=bold, italic=italic)
                i += 1
                continue
            inner = text[i + 1 : j]
            add_inline_runs(
                paragraph,
                inner,
                size_pt,
                font_name,
                bold=bold,
                italic=True,
            )
            i = j + 1
            continue

        start = i
        while i < n and text[i] not in "`*":
            i += 1
        if i > start:
            _emit_run(
                paragraph,
                text[start:i],
                size_pt,
                font_name,
                bold=bold,
                italic=italic,
            )


def is_table_row(line: str) -> bool:
    s = line.strip()
    return s.startswith("|") and s.count("|") >= 2


def is_table_sep(line: str) -> bool:
    s = line.strip().replace(" ", "")
    return bool(re.match(r"^\|?[-:|]+\|?$", s))


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    i = start
    while i < len(lines) and is_table_row(lines[i]):
        row = [c.strip() for c in lines[i].strip().strip("|").split("|")]
        if is_table_sep(lines[i]):
            i += 1
            continue
        rows.append(row)
        i += 1
    return rows, i


def is_caption_line(stripped: str) -> bool:
    inner = stripped.strip()
    if len(inner) >= 2 and inner.startswith("*") and inner.endswith("*"):
        inner = inner[1:-1].strip()
    else:
        return False
    return bool(re.match(r"^(Figura|Tabela)\s+\d+", inner, re.I))


def _warn_missing_figures(md_path: Path) -> None:
    md_dir = md_path.parent
    missing: list[Path] = []
    for line in md_path.read_text(encoding="utf-8").splitlines():
        m = re.match(r"!\[[^\]]*\]\(([^)]+)\)", line.strip())
        if not m:
            continue
        rel = m.group(1).strip()
        img = (md_dir / rel).resolve()
        if not img.is_file():
            missing.append(img)
    if not missing:
        return
    print(
        "Aviso: imagens referenciadas no .md em falta — o DOCX fica muito curto sem elas:",
        file=sys.stderr,
    )
    for p in missing:
        print(f"  - {p}", file=sys.stderr)
    print("Corra: python scripts/gerar_figuras_relatorio.py", file=sys.stderr)


def load_cabecalho_json(path: Path) -> dict | None:
    if not path.is_file():
        return None
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except (json.JSONDecodeError, OSError) as e:
        print("Aviso: não foi possível ler o cabeçalho JSON:", path, e, file=sys.stderr)
        return None


def _add_horizontal_rule_paragraph(doc: Document) -> None:
    """Linha horizontal fina (estilo capa IDP), cor azul aproximada."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(" ")
    run.font.size = Pt(2)
    p_pr = p._element.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    for edge in ("top", "left", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        p_bdr.append(el)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2F5496")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def _cover_label_line(doc: Document, label: str, value: str, font_name: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(label)
    r1.bold = True
    set_run_font(r1, 12, font_name)
    r2 = p.add_run(" " + (value or "").strip())
    set_run_font(r2, 12, font_name)


def add_cover_idp(doc: Document, cfg: dict, font_name: str) -> None:
    """
    Capa no formato IDP (campos administrativos).
    Chaves JSON opcionais:
    titulo_capa (ex.: PROJETO FINAL; se vazio ou omitido, não é impresso),
    disciplina, curso, area_concentracao, docente_responsavel, email_docente,
    ano_bimestre_referencia, entrega_professora (data/ texto da entrega),
    entrega_rotulo (opcional; por defeito «ENTREGA:»), texto_modelo_opcional,
    titulo_trabalho, nome_aluno
    """
    titulo = str(cfg.get("titulo_capa", "")).strip()
    if titulo:
        p0 = doc.add_paragraph()
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p0.paragraph_format.space_after = Pt(18)
        p0.paragraph_format.first_line_indent = Cm(0)
        r0 = p0.add_run(titulo.upper())
        r0.bold = True
        set_run_font(r0, 14, font_name)

    _cover_label_line(doc, "DISCIPLINA:", str(cfg.get("disciplina", "")), font_name)
    _cover_label_line(doc, "CURSO:", str(cfg.get("curso", "")), font_name)
    _cover_label_line(doc, "ÁREA DE CONCENTRAÇÃO:", str(cfg.get("area_concentracao", "")), font_name)
    _cover_label_line(doc, "DOCENTE RESPONSÁVEL:", str(cfg.get("docente_responsavel", "")), font_name)
    email_doc = str(cfg.get("email_docente", "")).strip()
    if email_doc:
        _cover_label_line(doc, "E-MAIL:", email_doc, font_name)
    _cover_label_line(doc, "ANO E BIMESTRE DE REFERÊNCIA:", str(cfg.get("ano_bimestre_referencia", "")), font_name)

    _add_horizontal_rule_paragraph(doc)

    rotulo_entrega = str(cfg.get("entrega_rotulo", "ENTREGA:")).strip()
    if rotulo_entrega and not rotulo_entrega.endswith(":"):
        rotulo_entrega += ":"
    _cover_label_line(doc, rotulo_entrega or "ENTREGA:", str(cfg.get("entrega_professora", "")), font_name)

    modelo = (cfg.get("texto_modelo_opcional") or "").strip()
    if modelo:
        pm = doc.add_paragraph()
        pm.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pm.paragraph_format.first_line_indent = Cm(0)
        pm.paragraph_format.space_before = Pt(10)
        pm.paragraph_format.space_after = Pt(4)
        add_inline_runs(pm, modelo, 12, font_name)

    tt = (cfg.get("titulo_trabalho") or "").strip()
    if tt:
        pt = doc.add_paragraph()
        pt.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pt.paragraph_format.first_line_indent = Cm(0)
        pt.paragraph_format.space_after = Pt(4)
        rt = pt.add_run("Título do trabalho: ")
        rt.bold = True
        set_run_font(rt, 12, font_name)
        add_inline_runs(pt, tt, 12, font_name)

    na = (cfg.get("nome_aluno") or "").strip()
    pn = doc.add_paragraph()
    pn.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pn.paragraph_format.first_line_indent = Cm(0)
    pn.paragraph_format.space_before = Pt(6)
    rn = pn.add_run("Nome do(a) aluno(a): ")
    rn.bold = True
    set_run_font(rn, 12, font_name)
    add_inline_runs(pn, na or "(Preencher o nome completo)", 12, font_name)


def export_main(
    md_path: Path,
    out_path: Path,
    font_name: str,
    margin_left: float,
    margin_other: float,
    cabecalho: dict | None = None,
) -> None:
    if not md_path.is_file():
        print("Arquivo não encontrado:", md_path, file=sys.stderr)
        sys.exit(1)

    _warn_missing_figures(md_path)

    raw_lines = md_path.read_text(encoding="utf-8").splitlines()
    lines = [ln.rstrip() for ln in raw_lines]

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = font_name
    style.font.size = Pt(12)

    section0 = doc.sections[0]
    apply_section_margins(section0, margin_left, margin_other)

    if cabecalho:
        add_cover_idp(doc, cabecalho, font_name)
        # Secção seguinte em fluxo contínuo (sem nova página): o corpo começa logo após
        # o último parágrafo da capa, na mesma página se houver espaço.
        body_section = doc.add_section(WD_SECTION_START.CONTINUOUS)
        apply_section_margins(body_section, margin_left, margin_other)
        _setup_section_footer_page_number(body_section, font_name, restart_from=1)
        _merge_empty_section_break_paragraph_into_previous(doc)
    else:
        _setup_section_footer_page_number(section0, font_name, restart_from=None)

    md_dir = md_path.parent
    i = 0
    in_references = False

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped == "---":
            i += 1
            continue

        if stripped.startswith("<!--"):
            while i < len(lines) and "-->" not in lines[i]:
                i += 1
            i += 1
            continue

        if stripped.startswith("*[") and "Inserir figura" in stripped:
            i += 1
            continue

        if stripped.startswith("![") and "](" in stripped:
            m = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", stripped)
            if m:
                rel = m.group(2).strip()
                img_path = (md_dir / rel).resolve()
                if img_path.is_file():
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.first_line_indent = Cm(0)
                    run = p.add_run()
                    run.add_picture(str(img_path), width=Cm(10))
                    p.paragraph_format.space_after = Pt(2)
                else:
                    p = doc.add_paragraph()
                    add_inline_runs(p, f"[Imagem não encontrada: {rel}]", 12, font_name)
                    format_paragraph_body(p, font_name, first_line=not in_references)
            i += 1
            continue

        if stripped.startswith("# "):
            t = stripped[2:].strip()
            p = doc.add_heading(t, level=1)
            format_heading_left(p, font_name, 14)
            i += 1
            continue

        if stripped.startswith("## "):
            t = stripped[3:].strip()
            in_references = bool(re.match(r"^(\d+\.\s*)?Referências\b", t, re.I))
            p = doc.add_heading(t, level=2)
            format_heading_left(p, font_name, 13)
            i += 1
            continue

        if stripped.startswith("### "):
            t = stripped[4:].strip()
            p = doc.add_heading(t, level=3)
            format_heading_left(p, font_name, 12)
            i += 1
            continue

        if stripped.startswith("> "):
            inner = stripped[2:].strip()
            p = doc.add_paragraph()
            add_inline_runs(p, inner, 11, font_name)
            format_paragraph_blockquote(p, font_name)
            i += 1
            continue

        if stripped.startswith("- ") or (
            stripped.startswith("* ") and len(stripped) > 2
        ):
            p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, stripped[2:].strip(), 12, font_name)
            format_list_item(p, font_name)
            i += 1
            continue

        if len(stripped) >= 2 and stripped.startswith("*") and stripped.endswith("*"):
            inner = stripped[1:-1]
            if inner and not inner.startswith("*"):
                if is_caption_line(stripped):
                    p = doc.add_paragraph()
                    r = p.add_run(inner)
                    r.italic = True
                    set_run_font(r, 12, font_name)
                    format_paragraph_caption(p, font_name)
                else:
                    p = doc.add_paragraph()
                    r = p.add_run(inner)
                    r.italic = True
                    set_run_font(r, 12, font_name)
                    if in_references:
                        format_paragraph_reference_entry(p, font_name)
                    else:
                        format_paragraph_body(p, font_name, first_line=True)
                i += 1
                continue

        if is_table_row(line):
            table_rows, new_i = parse_table(lines, i)
            if not table_rows:
                i = new_i
                continue
            cols = len(table_rows[0])
            table = doc.add_table(rows=len(table_rows), cols=cols)
            table.style = "Table Grid"
            for ri, row_cells in enumerate(table_rows):
                for ci, cell_text in enumerate(row_cells):
                    if ci < len(table.rows[ri].cells):
                        cell = table.rows[ri].cells[ci]
                        cell.text = strip_table_cell_markdown(cell_text)
                        for para in cell.paragraphs:
                            if in_references:
                                format_paragraph_reference_entry(para, font_name)
                            else:
                                format_paragraph_body(para, font_name, first_line=False)
            doc.add_paragraph()
            i = new_i
            continue

        p = doc.add_paragraph()
        add_inline_runs(p, stripped, 12, font_name)
        if in_references:
            format_paragraph_reference_entry(p, font_name)
        else:
            format_paragraph_body(p, font_name, first_line=True)
        i += 1

    out_path.parent.mkdir(parents=True, exist_ok=True)
    _enable_update_fields_on_open(doc)
    doc.save(out_path)
    print("Gerado:", out_path.resolve())


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Exporta o relatório Markdown para .docx (texto + capa IDP opcional via JSON; rodapé com n.º de página; sem sumário automático)."
    )
    parser.add_argument(
        "--font",
        choices=("times", "arial"),
        default="times",
        help="Fonte (por defeito: Times New Roman)",
    )
    parser.add_argument(
        "--margin-left",
        type=float,
        default=MARGIN_LEFT_CM,
        metavar="CM",
        help=f"Margem esquerda em cm (por defeito: {MARGIN_LEFT_CM})",
    )
    parser.add_argument(
        "--margin-other",
        type=float,
        default=MARGIN_OTHER_CM,
        metavar="CM",
        help=f"Margens superior, inferior e direita em cm (por defeito: {MARGIN_OTHER_CM})",
    )
    parser.add_argument(
        "-o",
        "--output",
        type=Path,
        default=OUT_PATH,
        help="Caminho do .docx de saída",
    )
    parser.add_argument(
        "-i",
        "--input",
        type=Path,
        default=MD_PATH,
        help="Caminho do .md de entrada",
    )
    parser.add_argument(
        "--cabecalho",
        type=Path,
        default=None,
        metavar="JSON",
        help="Ficheiro JSON da capa IDP (por defeito: relatorio_cabecalho.json na raiz, se existir)",
    )
    parser.add_argument(
        "--no-cabecalho",
        action="store_true",
        help="Não incluir capa IDP no DOCX",
    )
    parser.add_argument(
        "--gerar-figuras",
        action="store_true",
        help="Antes do export, executa scripts/gerar_figuras_relatorio.py (requer Parquet em data/)",
    )
    args = parser.parse_args()
    font = FONT_ARIAL if args.font == "arial" else FONT_TIMES
    if args.gerar_figuras:
        gen = ROOT / "scripts" / "gerar_figuras_relatorio.py"
        r = subprocess.run([sys.executable, str(gen)], cwd=str(ROOT))
        if r.returncode != 0:
            sys.exit(r.returncode)
    cab_path = args.cabecalho
    if cab_path is None and not args.no_cabecalho:
        cab_path = ROOT / "relatorio_cabecalho.json"
    cabecalho: dict | None = None
    if not args.no_cabecalho and cab_path is not None:
        cabecalho = load_cabecalho_json(cab_path)

    export_main(
        args.input,
        args.output,
        font_name=font,
        margin_left=args.margin_left,
        margin_other=args.margin_other,
        cabecalho=cabecalho,
    )


if __name__ == "__main__":
    main()
